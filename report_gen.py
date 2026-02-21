from docx import Document
from docx.shared import Inches
import datetime
import plotly.io as pio
import streamlit as st

def save_truss_plot(fig, filename):
    """
    Explicitly uses Kaleido to save the Plotly figure as a PNG.
    This bypasses the Chrome requirement on Streamlit Cloud.
    """
    try:
        # engine="kaleido" is the key fix for your GitHub/Streamlit deployment
        fig.write_image(filename, engine="kaleido", format="png", scale=3, width=1000, height=800)
        return True
    except Exception as e:
        st.error(f"Kaleido Export Error: {e}")
        return False

def generate_report(truss_system, fig_base=None, fig_res=None, scale_factor=1000.0, unit_label="kN"):
    doc = Document()
    image_base_path = "temp_base_plot.png"
    image_res_path = "temp_res_plot.png"
    
    # 1. Header & Software Details
    doc.add_heading('Structural Analysis Report', 0)
    doc.add_paragraph(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_heading('Software Specifications', level=1)
    soft_table = doc.add_table(rows=1, cols=2)
    soft_table.style = 'Table Grid'
    hdr_cells = soft_table.rows[0].cells
    hdr_cells[0].text = 'Property'
    hdr_cells[1].text = 'Details'
    
    software_info = [
        ('Software Name', 'Professional Truss Suite'),
        ('Developer', 'Mr. D Mandal, Assistant Professor, KITS Ramtek'),
        ('Core Engine', 'Direct Stiffness Method (DSM)'),
        ('Analysis Type', 'Linear Static 2D Truss Analysis'),
        ('Visualization', 'Plotly Rendering Engine')
    ]
    
    for prop, detail in software_info:
        row_cells = soft_table.add_row().cells
        row_cells[0].text = prop
        row_cells[1].text = detail

    # 2. Material Properties Table
    doc.add_heading('Material & Section Properties', level=1)
    mat_table = doc.add_table(rows=1, cols=3)
    mat_table.style = 'Table Grid'
    mat_table.rows[0].cells[0].text = 'Member ID'
    mat_table.rows[0].cells[1].text = 'Area (sq.m)'
    mat_table.rows[0].cells[2].text = 'E (N/sq.m)'
    
    for mbr in truss_system.members:
        row = mat_table.add_row().cells
        row[0].text, row[1].text, row[2].text = str(mbr.id), f"{mbr.A:.2e}", f"{mbr.E:.2e}"

    # 3. Image Section: Undeformed Geometry
    if fig_base:
        doc.add_heading('Undeformed Geometry Visualization', level=1)
        if save_truss_plot(fig_base, image_base_path):
            try:
                doc.add_picture(image_base_path, width=Inches(5.5))
            except:
                doc.add_paragraph("Error: Picture could not be embedded into the document.")
        else:
            doc.add_paragraph("Warning: Image generation failed due to environment constraints.")

    # 4. Image Section: Structural Forces (Results)
    if fig_res:
        doc.add_heading('Structural Forces Visualization', level=1)
        if save_truss_plot(fig_res, image_res_path):
            try:
                doc.add_picture(image_res_path, width=Inches(5.5))
            except:
                doc.add_paragraph("Error: Picture could not be embedded into the document.")
        else:
            doc.add_paragraph("Warning: Image generation failed due to environment constraints.")

    # 5. Nodal Displacements Table
    doc.add_heading('Nodal Displacements', level=1)
    disp_table = doc.add_table(rows=1, cols=3)
    disp_table.style = 'Table Grid'
    disp_table.rows[0].cells[0].text = 'Node ID'
    disp_table.rows[0].cells[1].text = 'Ux (m)'
    disp_table.rows[0].cells[2].text = 'Uy (m)'
    
    for n in truss_system.nodes:
        row = disp_table.add_row().cells
        row[0].text = str(n.id)
        row[1].text = f"{n.ux:.6e}"
        row[2].text = f"{n.uy:.6e}"

    # 5.5 Support Reactions Table (NEW ADDITION)
    doc.add_heading('Support Reactions', level=1)
    rxn_table = doc.add_table(rows=1, cols=3)
    rxn_table.style = 'Table Grid'
    rxn_table.rows[0].cells[0].text = 'Node ID'
    rxn_table.rows[0].cells[1].text = f'Rx ({unit_label})'
    rxn_table.rows[0].cells[2].text = f'Ry ({unit_label})'
    
    has_reactions = False
    for n in truss_system.nodes:
        if n.rx == 1 or n.ry == 1:
            has_reactions = True
            row = rxn_table.add_row().cells
            row[0].text = str(n.id)
            row[1].text = str(round(n.rx_val / scale_factor, 2)) if n.rx == 1 else "0.0"
            row[2].text = str(round(n.ry_val / scale_factor, 2)) if n.ry == 1 else "0.0"
            
    if not has_reactions:
        doc.add_paragraph("No rigid support reactions calculated.")
        
    

    # 6. Results Table (Dynamic Units)
    doc.add_heading('Detailed Analysis Results', level=1)
    res_table = doc.add_table(rows=1, cols=3)
    res_table.style = 'Table Grid'
    res_table.rows[0].cells[0].text = 'Member'
    res_table.rows[0].cells[1].text = f'Force ({unit_label})' # Uses dynamic unit label
    res_table.rows[0].cells[2].text = 'Nature'

    for mbr in truss_system.members:
        f = mbr.calculate_force()
        row = res_table.add_row().cells
        row[0].text = str(mbr.id)
        row[1].text = str(round(abs(f) / scale_factor, 2)) # Uses dynamic scale factor
        
        if abs(f) / scale_factor < 0.01:
            row[2].text = "Zero-Force"
        else:
            row[2].text = "Compressive" if f < 0 else "Tensile"

    doc.save("Analysis_Report.docx")

